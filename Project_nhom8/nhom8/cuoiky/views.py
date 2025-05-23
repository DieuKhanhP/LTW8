# inventory/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.views import generic
from django.contrib import messages
from .models import HangHoa  # Import model HangHoa của bạn
from django.urls import reverse_lazy
from django.contrib.auth.models import Group

from django.db.models import Q  # Để dùng OR trong query

# Giả sử TINH_TRANG_TONKHO_CHOICES được định nghĩa ở đây hoặc import từ models
# TINH_TRANG_TONKHO_CHOICES = [ ('CON_HANG', 'Còn hàng'), ('HET_HANG', 'Hết hàng'), ...]

from django.shortcuts import render
from django.db.models import Sum, Count
from django.utils import timezone
from .models import HangHoa, NhapKho, XuatKho, KiemKe , KhachHang # Import các model của bạn
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model
from .forms import NhaCungCapForm

User = get_user_model()

def is_warehouse_manager(user):
    return user.groups.filter(name='Quản lý kho').exists() or user.is_superuser



def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('tongquan')  # Chuyển hướng đến trang chính sau khi đăng nhập
        else:
            return render(request, 'login.html', {'error': 'Tên đăng nhập hoặc mật khẩu không đúng.'})
    return render(request, 'login.html')


from .forms import ProfileForm
from .models import CustomUser

from django.db.models import Q
from datetime import datetime

@login_required
def profile_view(request, username=None):
    """Hiển thị và cập nhật thông tin profile của người dùng."""

    # Nếu có username và người dùng hiện tại là quản lý kho, hiển thị profile của nhân viên đó
    if username and is_warehouse_manager(request.user):
        viewed_user = get_object_or_404(User, username=username)
    else:
        viewed_user = request.user

    if request.method == 'POST':
        # Chỉ cho phép cập nhật nếu đang xem profile của chính mình
        if viewed_user.username != request.user.username and not is_warehouse_manager(request.user):
            messages.error(request, 'Bạn không có quyền cập nhật thông tin của người dùng khác.')
            return redirect('profile')

        profile_form = ProfileForm(request.POST, instance=viewed_user)

        if profile_form.is_valid():
            user = profile_form.save(commit=False)
            # Đảm bảo lưu trường phone
            user.phone = profile_form.cleaned_data.get('phone', '')
            # Đảm bảo lưu trường position
            user.position = profile_form.cleaned_data.get('position', '')
            user.save()

            messages.success(request, 'Thông tin cá nhân đã được cập nhật thành công!')
            if username and is_warehouse_manager(request.user):
                return redirect('profile_detail', username=username)
            else:
                return redirect('profile')
        else:
            messages.error(request, 'Có lỗi xảy ra khi cập nhật thông tin. Vui lòng kiểm tra lại.')
    else:
        profile_form = ProfileForm(instance=viewed_user)

    # Lấy danh sách nhân viên kho nếu người dùng là quản lý kho
    warehouse_staff = None
    if is_warehouse_manager(request.user):
        warehouse_staff_group = Group.objects.filter(name='Nhân viên kho').first()
        if warehouse_staff_group:
            warehouse_staff = User.objects.filter(groups=warehouse_staff_group)
        else:
            # Nếu không có nhóm "Nhân viên kho", lấy tất cả người dùng trừ quản lý và admin
            manager_group = Group.objects.filter(name='Quản lý kho').first()
            if manager_group:
                warehouse_staff = User.objects.exclude(groups=manager_group).exclude(is_superuser=True)
            else:
                warehouse_staff = User.objects.exclude(is_superuser=True)

    context = {
        'profile_form': profile_form,
        'viewed_user': viewed_user,
        'is_manager': is_warehouse_manager(request.user),
        'warehouse_staff': warehouse_staff,
        'viewing_self': viewed_user.username == request.user.username
    }

    return render(request, 'profile.html', context)

def get_dashboard_context(request):
    now = timezone.now()
    years = range(now.year - 5, now.year + 1)

    month = request.GET.get('month')
    quarter = request.GET.get('quarter')
    year = request.GET.get('year')

    filters = Q()
    if year:
        filters &= Q(thoi_gian__year=year)
    if month:
        filters &= Q(thoi_gian__month=month)
        quarter = None
    elif quarter:
        start_month = (int(quarter) - 1) * 3 + 1
        end_month = start_month + 2
        filters &= Q(thoi_gian__month__gte=start_month, thoi_gian__month__lte=end_month)

    filtered_nhapkho = NhapKho.objects.filter(filters)
    filtered_xuatkho = XuatKho.objects.filter(filters)
    filtered_kiemke = KiemKe.objects.filter(filters)

    return {
        'years': years,
        'total_import': NhapKho.objects.filter(tinh_trang='DA_DUYET').filter(filters).count(),
        'total_export': XuatKho.objects.filter(tinh_trang='DA_DUYET').filter(filters).count(),
        'con_hang_count': HangHoa.objects.filter(tinh_trang='CON_HANG').count(),
        'gan_het_hang_count': HangHoa.objects.filter(tinh_trang='GAN_HET_HANG').count(),
        'het_hang_count': HangHoa.objects.filter(tinh_trang='HET_HANG').count()or 0,
        'tong_hang_hoa_count': HangHoa.objects.count() or 0,
        'phieu_nhap_count': filtered_nhapkho.count() or 0,
        'gia_tri_nhap': filtered_nhapkho.filter(tinh_trang='DA_DUYET').aggregate(Sum('tong'))['tong__sum'] or 0,
        'phieu_nhap_cho_duyet': NhapKho.objects.filter(tinh_trang='CHO_DUYET').filter(filters).count() or 0,
        'phieu_xuat_count': filtered_xuatkho.count() or 0,
        'gia_tri_xuat': filtered_xuatkho.filter(tinh_trang='DA_DUYET').aggregate(Sum('tong'))['tong__sum'] or 0,
        'phieu_xuat_cho_duyet': XuatKho.objects.filter(tinh_trang='CHO_DUYET').filter(filters).count() or 0,
        'phieu_kiem_ke_count': filtered_kiemke.count() or 0,
        'phieu_kiem_ke_cho_duyet': KiemKe.objects.filter(tinh_trang='CHO_DUYET').filter(filters).count() or 0,

    }
@login_required
def dashboard_view(request):
    """
    View chính hiển thị tổng quan kho – dữ liệu thống kê sẽ được xử lý trong get_dashboard_context.
    """
    context = get_dashboard_context(request)
    context['is_manager'] = is_warehouse_manager(request.user)
    return render(request, 'cuoiky/tongquan.html', context)



from docx import Document
from django.http import HttpResponse
from datetime import datetime
from django.utils.html import format_html
@login_required
def export_baocao_doc(request):
    if not request.user.groups.filter(name='Quản lý kho').exists():
        return JsonResponse({'success': False, 'error': 'Bạn không có quyền in báo cáo.'}, status=403)
    year = request.GET.get('year')
    month = request.GET.get('month')
    quarter = request.GET.get('quarter')

    # Gọi lại dashboard để lấy context
    fake_request = request
    fake_request.GET = request.GET.copy()
    context = get_dashboard_context(request)

    # Tạo tài liệu Word
    doc = Document()
    doc.add_heading('BÁO CÁO TỔNG QUAN KHO - VLXD Xuân Trường', 0)

    # Thời gian báo cáo
    thoi_gian = "Tất cả thời gian"
    if month:
        thoi_gian = f"Tháng {month} "
    elif quarter:
        thoi_gian = f"Quý {quarter} "
    if year:
        thoi_gian += f"Năm {year}"

    doc.add_paragraph(f"Thời gian báo cáo: {thoi_gian}", style='Intense Quote')

    # Tổng quan nhập - xuất
    doc.add_heading('I. Tổng quan nhập - xuất', level=1)
    doc.add_paragraph(f"- Tổng số lượng hàng hóa được nhập: {context.get('total_import')} đơn vị")
    doc.add_paragraph(f"- Tổng số lượng hàng hóa được xuất: {context.get('total_export')} đơn vị")

    # Tình trạng tồn kho
    doc.add_heading('II. Tình trạng tồn kho', level=1)
    doc.add_paragraph(f"- Số loại hàng hóa còn hàng: {context.get('con_hang_count')}")
    doc.add_paragraph(f"- Số loại hàng gần hết: {context.get('gan_het_hang_count')}")
    doc.add_paragraph(f"- Số loại hàng đã hết: {context.get('het_hang_count')}")
    doc.add_paragraph(f"- Tổng số loại hàng hóa hiện có: {context.get('tong_hang_hoa_count')}")

    # Nhập kho chi tiết
    doc.add_heading('III. Nhập kho', level=1)
    doc.add_paragraph(f"- Số phiếu nhập trong kỳ: {context.get('phieu_nhap_thang_count')}")
    doc.add_paragraph(f"- Tổng giá trị nhập: {context.get('gia_tri_nhap_thang'):,.0f} VND")
    doc.add_paragraph(f"- Phiếu nhập chờ duyệt: {context.get('phieu_nhap_cho_duyet_count')}")

    # Xuất kho chi tiết
    doc.add_heading('IV. Xuất kho', level=1)
    doc.add_paragraph(f"- Số phiếu xuất trong kỳ: {context.get('phieu_xuat_thang_count')}")
    doc.add_paragraph(f"- Tổng giá trị xuất: {context.get('gia_tri_xuat_thang'):,.0f} VND")
    doc.add_paragraph(f"- Phiếu xuất chờ duyệt: {context.get('phieu_xuat_cho_duyet_count')}")

    # Kiểm kê chi tiết
    doc.add_heading('V. Kiểm kê kho', level=1)
    doc.add_paragraph(f"- Số phiếu kiểm kê đã lập: {context.get('phieu_kiem_ke_thang_count')}")
    doc.add_paragraph(f"- Phiếu kiểm kê chờ duyệt: {context.get('phieu_kiem_ke_cho_duyet_count')}")

    # Kết luận
    doc.add_heading('VI. Ghi chú', level=1)
    doc.add_paragraph("Dữ liệu trong báo cáo chỉ tính các phiếu đã được duyệt và cập nhật đến thời điểm hiện tại.")

    # Trả file Word về
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f'bao_cao_tong_quan_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response
from django.http import JsonResponse
from django.contrib import messages

from django.http import JsonResponse
from django.contrib import messages
from django.views.decorators.http import require_POST

from django.http import JsonResponse
from django.contrib import messages
from django.views.decorators.http import require_POST
from django.shortcuts import get_object_or_404, redirect


@require_POST
@login_required
def hanghoa_delete(request, ma_hang):
    if not is_quan_ly_kho(request.user):
        return JsonResponse({
            'success': False,
            'error': 'Bạn không có quyền xóa hàng hóa.'
        }, status=403)

    hanghoa = get_object_or_404(HangHoa, ma_hang=ma_hang)

    try:
        hanghoa.delete()
        return JsonResponse({
            'success': True,
            'message': f'Đã xóa hàng hóa {ma_hang} thành công.'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Lỗi khi xóa: {str(e)}'
        }, status=500)

from .models import NHOM_HANG_CHOICES, TINH_TRANG_TONKHO_CHOICES
class HangHoaListView(generic.ListView):
    model = HangHoa
    template_name = 'cuoiky/hanghoa_list.html'  # Đổi tên template
    context_object_name = 'hanghoa_list'  # Đổi tên biến context
    paginate_by = 10  # Hoặc số lượng bạn muốn

    def get_queryset(self):
        queryset = super().get_queryset().order_by('-ma_hang')
        query_ma = self.request.GET.get('ma_hang')
        query_ten = self.request.GET.get('ten_hang')
        query_nhom = self.request.GET.get('nhom_hang')
        query_trangthai = self.request.GET.get('tinh_trang')

        if query_ma:
            # Tìm kiếm gần đúng hoặc chính xác tùy bạn chọn
            queryset = queryset.filter(ma_hang__icontains=query_ma)
        if query_ten:
            queryset = queryset.filter(ten_hang__icontains=query_ten)

        if query_nhom:
            queryset = queryset.filter(nhom_hang=query_nhom)
        if query_trangthai:
            queryset = queryset.filter(tinh_trang=query_trangthai)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Giữ lại các tham số tìm kiếm khi chuyển trang pagination
        search_params = self.request.GET.copy()
        if 'page' in search_params:
            del search_params['page']
        context['search_params'] = search_params.urlencode()
        context['request'] = self.request
        context['NHOM_HANG_CHOICES'] = NHOM_HANG_CHOICES
        context['TINH_TRANG_TONKHO_CHOICES'] = TINH_TRANG_TONKHO_CHOICES
        return context


from .forms import HangHoaForm


class HangHoaCreateView(generic.CreateView):
    model = HangHoa
    form_class = HangHoaForm
    template_name = 'cuoiky/hanghoa_form.html'
    success_url = reverse_lazy('hanghoa-list')

    def get_initial(self):
        initial = super().get_initial()
        # Tạo mã hàng tự động khi mở form
        last_item = HangHoa.objects.order_by('-ma_hang').first()
        if last_item and last_item.ma_hang.startswith("HH"):
            try:
                last_number = int(last_item.ma_hang.replace("HH", ""))
                new_number = last_number + 1
            except ValueError:
                new_number = 1
        else:
            new_number = 1
        initial['ma_hang'] = f"HH{new_number:04d}"
        return initial

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = 'THÊM HÀNG HÓA MỚI'
        return context

    def form_valid(self, form):
        messages.success(self.request, "Đã thêm hàng hóa mới thành công!")
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, "Có lỗi xảy ra khi thêm hàng hóa. Vui lòng kiểm tra lại thông tin.")
        return super().form_invalid(form)


class HangHoaUpdateView(generic.UpdateView):
    model = HangHoa
    form_class = HangHoaForm
    template_name = 'cuoiky/hanghoa_form.html'
    pk_url_kwarg = 'ma_hang'
    context_object_name = 'hanghoa'
    success_url = reverse_lazy('hanghoa-list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = f'Cập nhật: {self.object.ten_hang}'
        return context

    def form_valid(self, form):
        messages.success(self.request, f"Đã cập nhật hàng hóa {self.object.ten_hang} thành công!")
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, "Có lỗi xảy ra khi cập nhật hàng hóa. Vui lòng kiểm tra lại thông tin.")
        return super().form_invalid(form)


# Bạn có thể thêm View xem chi tiết nếu cần
class HangHoaDetailView(generic.DetailView):
    model = HangHoa
    template_name = 'cuoiky/hanghoa_detail.html'  # Tạo template này nếu cần
    pk_url_kwarg = 'ma_hang'
    context_object_name = 'hanghoa'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = f'CHI TIẾT THÔNG TIN HÀNG HÓA : {self.object.ten_hang}'
        return context


from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib import messages
from django.db import transaction
from django.http import HttpResponseRedirect
from django.contrib.auth.mixins import LoginRequiredMixin
from django.utils import timezone
from django.db.models import Q

from .models import NhapKho, ChiTietNhap, HangHoa
from .forms import NhapKhoForm, ChiTietNhapFormSet

from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.views.generic import ListView, DetailView, CreateView, UpdateView
from django.forms import inlineformset_factory
from django.contrib import messages
from django.db import transaction
from django.http import HttpResponseRedirect, JsonResponse
from django.utils import timezone
import openpyxl

from .models import NhapKho, ChiTietNhap, HangHoa
from .forms import NhapKhoForm, ChiTietNhapForm
from django.http import JsonResponse


# Thay đổi view xóa phiếu nhập kho
def xoa_nhapkho(request, ma_nhap):
    if not is_quan_ly_kho(request.user):
            return JsonResponse({
                'status': 'error',
                'message': 'Bạn không có quyền xóa phiếu nhập kho.'
            }, status=403)


    nhapkho = get_object_or_404(NhapKho, ma_nhap=ma_nhap)
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'Yêu cầu không hợp lệ.'}, status=400)

    try:
        with transaction.atomic():
            #  Nếu phiếu đã duyệt → cần trừ lại tồn kho
            if nhapkho.tinh_trang == 'DA_DUYET':
                for chitiet in nhapkho.chi_tiet_nhap.all():
                    hanghoa = chitiet.ma_hang
                    hanghoa.so_luong_he_thong -= chitiet.so_luong_nhap
                    hanghoa.save()

                    ton_entries = TonKhoTheoKho.objects.filter(hang_hoa=hanghoa, kho=nhapkho.kho)
                    for ton in ton_entries:
                        ton.so_luong -= chitiet.so_luong_nhap
                        if ton.so_luong < 0:
                            ton.so_luong = 0
                        ton.save()
                    print(
                        f"[-] Trừ {chitiet.so_luong_nhap} khỏi {hanghoa.ma_hang} → tồn kho mới: {hanghoa.so_luong_he_thong}")

            #  Xóa chi tiết trước, rồi xóa phiếu
            nhapkho.chi_tiet_nhap.all().delete()
            nhapkho.delete()
            print(f"Đã xóa phiếu nhập kho {ma_nhap} và cập nhật tồn kho")

        return JsonResponse({
            'status': 'success',
            'message': f'Phiếu nhập kho {ma_nhap} và hàng hóa liên quan đã được xóa thành công.'
        })

    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Có lỗi xảy ra khi xóa phiếu nhập kho: {str(e)}'
        }, status=500)


# Danh sách phiếu nhập kho
class NhapKhoListView(ListView):
    model = NhapKho
    template_name = 'cuoiky/nhapkho_list.html'
    context_object_name = 'nhapkho_list'
    paginate_by = 10

    def get_queryset(self):
        queryset = NhapKho.objects.all()

        # Lọc theo mã phiếu
        ma_nhap = self.request.GET.get('ma_nhap')
        if ma_nhap:
            queryset = queryset.filter(ma_nhap__icontains=ma_nhap)

        # Lọc theo nguồn nhập
        nguon_nhap = self.request.GET.get('nguon_nhap')
        if nguon_nhap:
            queryset = queryset.filter(nguon_nhap__icontains=nguon_nhap)

        # Lọc theo tình trạng
        tinh_trang = self.request.GET.get('tinh_trang')
        if tinh_trang:
            queryset = queryset.filter(tinh_trang=tinh_trang)

        # Lọc theo khoảng thời gian
        tu_ngay = self.request.GET.get('tu_ngay')
        den_ngay = self.request.GET.get('den_ngay')
        if tu_ngay:
            queryset = queryset.filter(thoi_gian__gte=tu_ngay)
        if den_ngay:
            queryset = queryset.filter(thoi_gian__lte=den_ngay)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Thêm các tham số tìm kiếm vào context để giữ lại khi phân trang
        search_params = ''
        for key, value in self.request.GET.items():
            if key != 'page':
                search_params += f'&{key}={value}'
        context['search_params'] = search_params
        return context


# Xem chi tiết phiếu nhập kho
class NhapKhoDetailView(DetailView):
    model = NhapKho
    template_name = 'cuoiky/nhapkho_detail.html'
    context_object_name = 'nhapkho'
    pk_url_kwarg = 'ma_nhap'

    def get_object(self, queryset=None):
        obj = get_object_or_404(NhapKho, ma_nhap=self.kwargs.get('ma_nhap'))
        # Đảm bảo tổng tiền được cập nhật trước khi hiển thị
        obj.tinh_lai_tong_tien()
        return obj


# Tạo phiếu nhập kho mới

class NhapKhoCreateView(CreateView):
    model = NhapKho
    form_class = NhapKhoForm
    template_name = 'cuoiky/nhapkho_form.html'
    success_url = reverse_lazy('nhapkho-list')

    def get_initial(self):
        initial = super().get_initial()
        # Tạo mã nhập kho tự động khi mở form
        prefix = 'NK'
        last_nhap = NhapKho.objects.order_by('-ma_nhap').first()
        if last_nhap and last_nhap.ma_nhap.startswith(prefix):
            try:
                last_number = int(last_nhap.ma_nhap.replace(prefix, ''))
                new_number = last_number + 1
            except ValueError:
                new_number = 1
        else:
            new_number = 1
        initial['ma_nhap'] = f"{prefix}{new_number:04d}"
        return initial

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = 'THÊM MỚI PHIẾU NHẬP KHO'
        if self.request.POST:
            context['chitiet_formset'] = ChiTietNhapFormSet(self.request.POST)
        else:
            context['chitiet_formset'] = ChiTietNhapFormSet(queryset=ChiTietNhap.objects.none())

        # Lấy tất cả hàng hóa với đầy đủ thông tin bao gồm đơn giá nhập
        context['all_hanghoa'] = HangHoa.objects.all().values('ma_hang', 'ten_hang', 'don_vi_tinh', 'don_gia_nhap')
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        chitiet_formset = context['chitiet_formset']


        # Lưu phiếu nhập kho trước
        with transaction.atomic():
            try:
                # Lưu phiếu nhập kho
                self.object = form.save(commit=False)
                self.object.tao_boi = self.request.user if self.request.user.is_authenticated else None
                self.object.save()
                print(f"Đã lưu phiếu nhập kho: {self.object.ma_nhap}")

                # Kiểm tra xem có ít nhất một chi tiết hợp lệ không
                has_valid_details = False
                valid_forms = []

                # In ra dữ liệu POST để kiểm tra
                print("POST data:", self.request.POST)

                # Kiểm tra tính hợp lệ của formset
                if not chitiet_formset.is_valid():
                    print("Formset không hợp lệ:", chitiet_formset.errors)
                    for i, form_errors in enumerate(chitiet_formset.errors):
                        print(f"Form {i} errors:", form_errors)
                    if chitiet_formset.non_form_errors():
                        print("Non-form errors:", chitiet_formset.non_form_errors())
                    messages.error(self.request, "Có lỗi xảy ra khi xác thực chi tiết nhập kho")
                    return self.form_invalid(form)

                # Xử lý từng form trong formset
                for i, detail_form in enumerate(chitiet_formset.forms):
                    print(f"Xử lý form {i}:")
                    print(f"  Form is valid: {detail_form.is_valid()}")
                    print(f"  Form has changed: {detail_form.has_changed()}")
                    print(f"  Form changed data: {detail_form.changed_data}")

                    # Chỉ xử lý các form hợp lệ, có dữ liệu và không bị đánh dấu xóa
                    if detail_form.is_valid() and detail_form.has_changed():
                        if not detail_form.cleaned_data.get('DELETE', False):
                            if detail_form.cleaned_data.get('ma_hang'):
                                has_valid_details = True
                                valid_forms.append(detail_form)
                                print(f"  Form {i} hợp lệ và sẽ được lưu")
                            else:
                                print(f"  Form {i} không có mã hàng")
                        else:
                            print(f"  Form {i} bị đánh dấu xóa")
                    else:
                        print(f"  Form {i} không hợp lệ hoặc không có thay đổi")

                if not has_valid_details:
                    messages.error(self.request, "Vui lòng thêm ít nhất một hàng hóa vào phiếu nhập kho.")
                    # Xóa phiếu nhập kho đã tạo vì không có chi tiết hợp lệ
                    self.object.delete()
                    return self.form_invalid(form)

                # Lưu từng chi tiết nhập kho riêng biệt
                for detail_form in valid_forms:
                    try:
                        chitiet = detail_form.save(commit=False)
                        chitiet.ma_nhap = self.object

                        # Đảm bảo các trường không bị null
                        if chitiet.don_gia_nhap is None:
                            if chitiet.ma_hang and chitiet.ma_hang.don_gia_nhap:
                                chitiet.don_gia_nhap = chitiet.ma_hang.don_gia_nhap
                            else:
                                chitiet.don_gia_nhap = 0

                        if chitiet.chiet_khau is None:
                            chitiet.chiet_khau = 0

                        # Tính thanh tiền
                        chitiet.thanh_tien = (chitiet.don_gia_nhap * chitiet.so_luong_nhap) * (
                                1 - chitiet.chiet_khau / 100)

                        # Lưu chi tiết
                        chitiet.save()
                        print(f"Đã lưu chi tiết nhập kho: {chitiet}")
                    except Exception as e:
                        print(f"Lỗi khi lưu chi tiết: {e}")
                        import traceback
                        traceback.print_exc()

                # Cập nhật tổng tiền
                self.object.tinh_lai_tong_tien()
                print(f"Đã cập nhật tổng tiền: {self.object.tong}")

                messages.success(self.request, "Tạo phiếu nhập kho thành công!")
                return redirect('nhapkho-detail', ma_nhap=self.object.ma_nhap)
            except Exception as e:
                print(f"Lỗi khi lưu phiếu nhập kho: {e}")
                import traceback
                traceback.print_exc()
                messages.error(self.request, f"Có lỗi xảy ra khi lưu phiếu nhập kho: {e}")
                return self.form_invalid(form)

    def form_invalid(self, form):
        print("Lỗi trong form:", form.errors)
        return self.render_to_response(self.get_context_data(form=form))

    def get_success_url(self):
        return reverse_lazy('nhapkho-detail', kwargs={'ma_nhap': self.object.ma_nhap})


from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import HttpResponseForbidden


# Hàm kiểm tra người dùng có thuộc nhóm "Quản lý kho"
def is_quan_ly_kho(user):
    return user.groups.filter(name='Quản lý kho').exists() or user.is_superuser


# Cập nhật phiếu nhập kho
class NhapKhoUpdateView(UpdateView):
    model = NhapKho
    form_class = NhapKhoForm
    template_name = 'cuoiky/nhapkho_form.html'
    pk_url_kwarg = 'ma_nhap'
    success_url = reverse_lazy('nhapkho-list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = f"Cập nhật phiếu nhập kho: {self.object.ma_nhap}"
        if self.request.POST:
            context['chitiet_formset'] = ChiTietNhapFormSet(self.request.POST, instance=self.object)
        else:
            context['chitiet_formset'] = ChiTietNhapFormSet(instance=self.object)
        # Đảm bảo truyền đầy đủ thông tin hàng hóa bao gồm đơn giá nhập
        context['all_hanghoa'] = HangHoa.objects.all().values('ma_hang', 'ten_hang', 'don_vi_tinh', 'don_gia_nhap')
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        chitiet_formset = context['chitiet_formset']

        # Kiểm tra tính hợp lệ của formset trước khi lưu
        if not chitiet_formset.is_valid():
            print("Lỗi trong chi tiết nhập kho:", chitiet_formset.errors)
            for i, err in enumerate(chitiet_formset.errors):
                if err:
                    print(f"Form {i} errors:", err)
            for err in chitiet_formset.non_form_errors():
                print(f"Non-form error: {err}")
            messages.error(self.request, "Có lỗi xảy ra khi lưu chi tiết nhập kho")
            return self.form_invalid(form)

        with transaction.atomic():
            try:
                # Lưu phiếu nhập kho trước
                self.object = form.save(commit=False)
                self.object.tao_boi = self.request.user if self.request.user.is_authenticated else None
                self.object.save()
                print("Phiếu nhập kho đã được lưu:", self.object)

                # Gán instance cho formset và lưu
                chitiet_formset.instance = self.object
                saved_instances = chitiet_formset.save(commit=False)

                # Lưu từng chi tiết nhập kho
                for instance in saved_instances:
                    # Đảm bảo don_gia_nhap không null
                    if instance.don_gia_nhap is None:
                        if instance.ma_hang and instance.ma_hang.don_gia_nhap:
                            instance.don_gia_nhap = instance.ma_hang.don_gia_nhap
                        else:
                            instance.don_gia_nhap = 0

                    # Tính thanh_tien
                    instance.thanh_tien = (instance.don_gia_nhap * instance.so_luong_nhap) * (
                            1 - instance.chiet_khau / 100)
                    instance.save()
                    print(f"Chi tiết nhập kho đã được lưu: {instance}")

                # Lưu các quan hệ many-to-many nếu có
                chitiet_formset.save_m2m()

                # Cập nhật tổng tiền
                self.object.tinh_lai_tong_tien()
                print("Tổng tiền đã được cập nhật:", self.object.tong)

                messages.success(self.request, f"Phiếu nhập kho {self.object.ma_nhap} đã được cập nhật thành công.")
                return HttpResponseRedirect(self.get_success_url())
            except Exception as e:
                print(f"Lỗi khi lưu phiếu nhập kho: {e}")
                messages.error(self.request, f"Có lỗi xảy ra khi lưu phiếu nhập kho: {e}")
                return self.form_invalid(form)

    def form_invalid(self, form):
        print("Lỗi trong form:", form.errors)
        return self.render_to_response(self.get_context_data(form=form))

    def get_success_url(self):
        return reverse_lazy('nhapkho-detail', kwargs={'ma_nhap': self.object.ma_nhap})


from .models import TonKhoTheoKho, Kho
# Duyệt phiếu nhập kho
from django.db.models import Sum
from .models import NhapKho, TonKhoTheoKho, Kho


@login_required
def duyet_nhapkho(request, ma_nhap):
    if not is_quan_ly_kho(request.user):
        messages.error(request, "Bạn không có quyền duyệt phiếu nhập kho.")
        return redirect('nhapkho-detail', ma_nhap=ma_nhap)

    nhapkho = get_object_or_404(NhapKho, ma_nhap=ma_nhap)

    if nhapkho.tinh_trang != 'CHO_DUYET':
        messages.error(request, "Chỉ có thể duyệt phiếu đang ở trạng thái chờ duyệt.")
        return redirect('nhapkho-detail', ma_nhap=ma_nhap)

    kho_chinh = nhapkho.kho
    kho_du_phong = Kho.objects.get(ma_kho='KHO0002')  # Kho dự trữ

    ton_chinh = TonKhoTheoKho.objects.filter(kho=kho_chinh).aggregate(tong=Sum('so_luong'))['tong'] or 0
    ton_phu = TonKhoTheoKho.objects.filter(kho=kho_du_phong).aggregate(tong=Sum('so_luong'))['tong'] or 0

    gioi_han_chinh = kho_chinh.gioi_han
    gioi_han_phu = kho_du_phong.gioi_han

    cho_phep_chinh = gioi_han_chinh - ton_chinh
    cho_phep_phu = gioi_han_phu - ton_phu
    tong_cho_phep = cho_phep_chinh + cho_phep_phu

    tong_nhap = sum(ct.so_luong_nhap for ct in nhapkho.chi_tiet_nhap.all())

    if tong_nhap > tong_cho_phep:
        messages.error(request,
                       f"Không thể duyệt. Tổng hàng ({tong_nhap}) vượt quá khả năng chứa của cả hai kho ({tong_cho_phep}).")
        return redirect('nhapkho-detail', ma_nhap=ma_nhap)

    with transaction.atomic():
        nhapkho.tinh_trang = 'DA_DUYET'
        nhapkho.save()

        for chitiet in nhapkho.chi_tiet_nhap.all():
            hh = chitiet.ma_hang
            so_luong_can_nhap = chitiet.so_luong_nhap

            # Nhập vào kho chính trước
            nhap_chinh = min(so_luong_can_nhap, cho_phep_chinh)
            if nhap_chinh > 0:
                ton_kho_chinh, _ = TonKhoTheoKho.objects.get_or_create(
                    kho=kho_chinh,
                    hang_hoa=hh,
                    defaults={'so_luong': 0}
                )
                ton_kho_chinh.so_luong += nhap_chinh
                ton_kho_chinh.save()
                hh.so_luong_he_thong += nhap_chinh
                hh.save()
                cho_phep_chinh -= nhap_chinh

            # Nhập phần dư vào kho phụ
            con_lai = so_luong_can_nhap - nhap_chinh
            if con_lai > 0:
                nhap_phu = min(con_lai, cho_phep_phu)
                if nhap_phu > 0:
                    ton_kho_phu, _ = TonKhoTheoKho.objects.get_or_create(
                        kho=kho_du_phong,
                        hang_hoa=hh,
                        defaults={'so_luong': 0}
                    )
                    ton_kho_phu.so_luong += nhap_phu
                    ton_kho_phu.save()
                    hh.so_luong_he_thong += nhap_phu
                    hh.save()
                    cho_phep_phu -= nhap_phu

    messages.success(request, f"Phiếu nhập kho {ma_nhap} đã được duyệt và cập nhật tồn kho.")
    return redirect('nhapkho-detail', ma_nhap=ma_nhap)


# Từ chối phiếu nhập kho
def tu_choi_nhapkho(request, ma_nhap):
    if not is_quan_ly_kho(request.user):
        messages.error(request, "Bạn không có quyền duyệt phiếu nhập kho.")
        return redirect('nhapkho-detail', ma_nhap=ma_nhap)
    nhapkho = get_object_or_404(NhapKho, ma_nhap=ma_nhap)
    print(f"Từ chối phiếu nhập kho: {ma_nhap}")

    if nhapkho.tinh_trang == 'CHO_DUYET':
        nhapkho.tinh_trang = 'TU_CHOI'
        nhapkho.save()
        print(f"Đã từ chối phiếu nhập kho {ma_nhap}")
        messages.success(request, f"Phiếu nhập kho {ma_nhap} đã bị từ chối.")
    else:
        print(f"Không thể từ chối phiếu {ma_nhap} vì trạng thái là {nhapkho.tinh_trang}")
        messages.error(request, "Chỉ có thể từ chối phiếu đang ở trạng thái chờ duyệt.")

    return redirect('nhapkho-detail', ma_nhap=ma_nhap)


def import_hanghoa_excel(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        try:
            workbook = openpyxl.load_workbook(excel_file)
            sheet = workbook.active
            created_count = 0
            updated_count = 0
            errors = []

            # Giả định dòng đầu tiên là header
            header = [cell.value for cell in sheet[1]]
            expected_headers = ['ma_hang', 'ten_hang', 'nhom_hang', 'don_vi_tinh', 'so_luong_he_thong', 'han_su_dung',
                                'mo_ta', 'tinh_trang', 'url_image']

            if header != expected_headers:
                messages.error(request,
                               f"Cấu trúc file Excel không đúng. Mong đợi các cột: {', '.join(expected_headers)}")
                return redirect('nhapkho-create')  # Hoặc trang nào bạn muốn

            for row_index, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
                row_data = dict(zip(header, row))
                ma_hang = row_data.get('ma_hang')
                print(f"Đang xử lý dòng {row_index} với dữ liệu: {row_data}")

                if not ma_hang:
                    errors.append(f"Dòng {row_index}: Mã hàng không được để trống.")
                    continue

                try:
                    hang_hoa, created = HangHoa.objects.update_or_create(
                        ma_hang=ma_hang,
                        defaults={
                            'ten_hang': row_data.get('ten_hang', ''),
                            'nhom_hang': row_data.get('nhom_hang', ''),
                            'don_vi_tinh': row_data.get('don_vi_tinh', ''),
                            'so_luong_he_thong': row_data.get('so_luong_he_thong', 0),
                            'han_su_dung': row_data.get('han_su_dung'),
                            'mo_ta': row_data.get('mo_ta', ''),
                            'tinh_trang': row_data.get('tinh_trang', True),  # Giả định True nếu không có
                            'url_image': row_data.get('url_image', '')  # Cần xử lý file nếu có
                        }
                    )
                    if created:
                        created_count += 1
                        print(f"Đã tạo mới hàng hóa: {hang_hoa}")
                    else:
                        updated_count += 1
                        print(f"Đã cập nhật hàng hóa: {hang_hoa}")
                except Exception as e:
                    error_message = f"Dòng {row_index}: Lỗi khi xử lý - {e}"
                    errors.append(error_message)
                    print(error_message)

            if created_count > 0:
                messages.success(request, f"Đã thêm mới {created_count} hàng hóa từ Excel.")
            if updated_count > 0:
                messages.info(request, f"Đã cập nhật {updated_count} hàng hóa từ Excel.")  # changed to info
            if errors:
                messages.error(request, "Có lỗi xảy ra trong quá trình nhập liệu từ Excel:")
                for error_message in errors:  # changed variable name to avoid shadowing
                    messages.error(request, error_message)

        except Exception as e:
            error_message = f"Lỗi khi đọc file Excel: {e}"
            messages.error(request, error_message)
            print(error_message)

    return redirect('nhapkho-create')  # Chuyển hướng về trang tạo phiếu nhập kho


def get_hanghoa_info(request, hanghoa_id):
    try:
        hang_hoa = HangHoa.objects.get(pk=hanghoa_id)
        data = {
            'ma_hang': hang_hoa.ma_hang,
            'ten_hang': hang_hoa.ten_hang,
            'don_gia_nhap': float(hang_hoa.don_gia_nhap) if hang_hoa.don_gia_nhap else 0,
            'don_gia_ban': float(hang_hoa.don_gia_ban) if hang_hoa.don_gia_ban else 0,
            'so_luong_he_thong': hang_hoa.so_luong_he_thong,
            'don_vi_tinh': hang_hoa.don_vi_tinh
        }
        return JsonResponse(data)
    except HangHoa.DoesNotExist:
        return JsonResponse({'error': 'Không tìm thấy hàng hóa'}, status=404)


from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib import messages
from django.db import transaction
from django.http import HttpResponseRedirect
from django.contrib.auth.mixins import LoginRequiredMixin
from django.utils import timezone
from django.db.models import Q

from .models import XuatKho, ChiTietXuat, HangHoa
from .forms import XuatKhoForm, ChiTietXuatFormSet

# Xóa phiếu xuất kho
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.db import transaction
from .models import XuatKho


def xoa_xuatkho(request, ma_xuat):
    if not is_quan_ly_kho(request.user):
            return JsonResponse({
                'status': 'error',
                'message': 'Bạn không có quyền xóa phiếu xuất kho.'
            }, status=403)

    xuatkho = get_object_or_404(XuatKho, ma_xuat=ma_xuat)

    try:
        with transaction.atomic():
            if xuatkho.tinh_trang == 'DA_DUYET':
                for chitiet in xuatkho.chi_tiet_xuat.all():
                    hanghoa = chitiet.ma_hang
                    hanghoa.so_luong_he_thong += chitiet.so_luong_xuat
                    hanghoa.save()

                    ton_kho = TonKhoTheoKho.objects.filter(hang_hoa=hanghoa, kho=xuatkho.kho).first()
                    if ton_kho:
                        ton_kho.so_luong += chitiet.so_luong_xuat
                        ton_kho.save()

            xuatkho.chi_tiet_xuat.all().delete()
            xuatkho.delete()
        return JsonResponse({'status': 'success', 'message': f'Đã xóa phiếu xuất kho {ma_xuat} thành công.'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Lỗi khi xóa: {str(e)}'}, status=500)



# Danh sách phiếu xuất kho
class XuatKhoListView(ListView):
    model = XuatKho
    template_name = 'cuoiky/xuatkho_list.html'
    context_object_name = 'xuatkho_list'
    paginate_by = 10

    def get_queryset(self):
        queryset = XuatKho.objects.all()

        # Lọc theo mã phiếu
        ma_xuat = self.request.GET.get('ma_xuat')
        if ma_xuat:
            queryset = queryset.filter(ma_xuat__icontains=ma_xuat)

        # Lọc theo nguồn xuất
        nguon_xuat = self.request.GET.get('nguon_xuat')
        if nguon_xuat:
            queryset = queryset.filter(nguon_xuat__icontains=nguon_xuat)

        # Lọc theo tình trạng
        tinh_trang = self.request.GET.get('tinh_trang')
        if tinh_trang:
            queryset = queryset.filter(tinh_trang=tinh_trang)

        # Lọc theo khoảng thời gian
        tu_ngay = self.request.GET.get('tu_ngay')
        den_ngay = self.request.GET.get('den_ngay')
        if tu_ngay:
            queryset = queryset.filter(thoi_gian__gte=tu_ngay)
        if den_ngay:
            queryset = queryset.filter(thoi_gian__lte=den_ngay)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Thêm các tham số tìm kiếm vào context để giữ lại khi phân trang
        search_params = ''
        for key, value in self.request.GET.items():
            if key != 'page':
                search_params += f'&{key}={value}'
        context['search_params'] = search_params
        return context


# Xem chi tiết phiếu xuất kho
class XuatKhoDetailView(DetailView):
    model = XuatKho
    template_name = 'cuoiky/xuatkho_detail.html'
    context_object_name = 'xuatkho'
    pk_url_kwarg = 'ma_xuat'

    def get_object(self, queryset=None):
        return get_object_or_404(XuatKho, ma_xuat=self.kwargs.get('ma_xuat'))


# Tạo phiếu xuất kho mới
class XuatKhoCreateView(CreateView):
    model = XuatKho
    form_class = XuatKhoForm
    template_name = 'cuoiky/xuatkho_form.html'

    def get_initial(self):
        initial = super().get_initial()
        # Tạo mã xuất kho tự động khi mở form
        prefix = 'XK'
        last_xuat = XuatKho.objects.order_by('-ma_xuat').first()
        if last_xuat and last_xuat.ma_xuat.startswith(prefix):
            try:
                last_number = int(last_xuat.ma_xuat.replace(prefix, ''))
                new_number = last_number + 1
            except ValueError:
                new_number = 1
        else:
            new_number = 1
        initial['ma_xuat'] = f"{prefix}{new_number:04d}"
        return initial

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = 'THÊM MỚI PHIẾU XUẤT KHO'
        if self.request.POST:
            context['chitiet_formset'] = ChiTietXuatFormSet(self.request.POST)
        else:
            context['chitiet_formset'] = ChiTietXuatFormSet()
        context['all_hanghoa'] = HangHoa.objects.filter(so_luong_he_thong__gt=0)
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        chitiet_formset = context['chitiet_formset']

        with transaction.atomic():
            try:
                # Lưu phiếu xuất kho
                self.object = form.save(commit=False)
                self.object.tao_boi = self.request.user if self.request.user.is_authenticated else None
                self.object.save()
                print(f"Đã lưu phiếu xuất kho: {self.object.ma_xuat}")

                # Kiểm tra xem có ít nhất một chi tiết hợp lệ không
                has_valid_details = False
                valid_forms = []

                # Kiểm tra tính hợp lệ của formset
                if not chitiet_formset.is_valid():
                    print("Formset không hợp lệ:", chitiet_formset.errors)
                    for i, form_errors in enumerate(chitiet_formset.errors):
                        print(f"Form {i} errors:", form_errors)
                    if chitiet_formset.non_form_errors():
                        print("Non-form errors:", chitiet_formset.non_form_errors())
                    messages.error(self.request, "Có lỗi xảy ra khi xác thực chi tiết xuất kho")
                    return self.form_invalid(form)

                # Xử lý từng form trong formset
                for i, detail_form in enumerate(chitiet_formset.forms):
                    # Chỉ xử lý các form hợp lệ, có dữ liệu và không bị đánh dấu xóa
                    if detail_form.is_valid() and detail_form.has_changed():
                        if not detail_form.cleaned_data.get('DELETE', False):
                            if detail_form.cleaned_data.get('ma_hang'):
                                has_valid_details = True
                                valid_forms.append(detail_form)
                            else:
                                print(f"Form {i} không có mã hàng")
                        else:
                            print(f"Form {i} bị đánh dấu xóa")
                    else:
                        print(f"Form {i} không hợp lệ hoặc không có thay đổi")

                if not has_valid_details:
                    messages.error(self.request, "Vui lòng thêm ít nhất một hàng hóa vào phiếu xuất kho.")
                    # Xóa phiếu xuất kho đã tạo vì không có chi tiết hợp lệ
                    self.object.delete()
                    return self.form_invalid(form)

                # Lưu từng chi tiết xuất kho riêng biệt
                for detail_form in valid_forms:
                    try:
                        chitiet = detail_form.save(commit=False)
                        chitiet.ma_xuat = self.object

                        # Đảm bảo các trường không bị null
                        if chitiet.don_gia_xuat is None:
                            if chitiet.ma_hang and chitiet.ma_hang.don_gia_ban:
                                chitiet.don_gia_xuat = chitiet.ma_hang.don_gia_ban
                            else:
                                chitiet.don_gia_xuat = 0

                        if chitiet.chiet_khau is None:
                            chitiet.chiet_khau = 0

                        # Tính thanh tiền
                        chitiet.thanh_tien = (chitiet.don_gia_xuat * chitiet.so_luong_xuat) * (
                                    1 - chitiet.chiet_khau / 100)

                        # Lưu chi tiết
                        chitiet.save()
                        print(f"Đã lưu chi tiết xuất kho: {chitiet}")
                    except Exception as e:
                        print(f"Lỗi khi lưu chi tiết: {e}")
                        import traceback
                        traceback.print_exc()

                # Cập nhật tổng tiền
                self.object.tinh_lai_tong_tien()
                print(f"Đã cập nhật tổng tiền: {self.object.tong}")

                messages.success(self.request, "Tạo phiếu xuất kho thành công!")
                return redirect('xuatkho-detail', ma_xuat=self.object.ma_xuat)
            except Exception as e:
                print(f"Lỗi khi lưu phiếu xuất kho: {e}")
                import traceback
                traceback.print_exc()
                messages.error(self.request, f"Có lỗi xảy ra khi lưu phiếu xuất kho: {e}")
                return self.form_invalid(form)

    def form_invalid(self, form):
        print("Lỗi trong form:", form.errors)
        return self.render_to_response(self.get_context_data(form=form))

    def get_success_url(self):
        return reverse_lazy('xuatkho-detail', kwargs={'ma_xuat': self.object.ma_xuat})


# Cập nhật phiếu xuất kho
class XuatKhoUpdateView(UpdateView):
    model = XuatKho
    form_class = XuatKhoForm
    template_name = 'cuoiky/xuatkho_form.html'
    pk_url_kwarg = 'ma_xuat'

    def get_object(self, queryset=None):
        return get_object_or_404(XuatKho, ma_xuat=self.kwargs.get('ma_xuat'))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.POST:
            context['chitiet_formset'] = ChiTietXuatFormSet(self.request.POST, instance=self.object)
        else:
            context['chitiet_formset'] = ChiTietXuatFormSet(instance=self.object)
        context['all_hanghoa'] = HangHoa.objects.filter(so_luong_he_thong__gt=0)
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        chitiet_formset = context['chitiet_formset']

        with transaction.atomic():
            self.object = form.save()

            if chitiet_formset.is_valid():
                chitiet_formset.instance = self.object
                chitiet_formset.save()

                # Cập nhật tổng tiền
                self.object.tinh_lai_tong_tien()

                messages.success(self.request, f"Phiếu xuất kho {self.object.ma_xuat} đã được cập nhật thành công.")
                return HttpResponseRedirect(self.get_success_url())
            else:
                return self.form_invalid(form)

    def get_success_url(self):
        return reverse_lazy('xuatkho-detail', kwargs={'ma_xuat': self.object.ma_xuat})


# Duyệt phiếu xuất kho
def duyet_xuatkho(request, ma_xuat):
    if not is_quan_ly_kho(request.user):
        messages.error(request, "Bạn không có quyền duyệt phiếu xuất kho.")
        return redirect('xuatkho-detail', ma_xuat=ma_xuat)

    xuatkho = get_object_or_404(XuatKho, ma_xuat=ma_xuat)

    if xuatkho.tinh_trang != 'CHO_DUYET':
        messages.error(request, "Chỉ có thể duyệt phiếu đang ở trạng thái chờ duyệt.")
        return redirect('xuatkho-detail', ma_xuat=ma_xuat)

    # Kiểm tra tồn kho trước khi duyệt
    du_ton_kho = True
    thieu_hang = []

    with transaction.atomic():
        for chitiet in xuatkho.chi_tiet_xuat.all():
            hang_hoa = chitiet.ma_hang
            so_luong_xuat = chitiet.so_luong_xuat

            # Kiểm tra tồn kho
            if hang_hoa.so_luong_he_thong < so_luong_xuat:
                du_ton_kho = False
                thieu_hang.append({
                    'ten_hang': hang_hoa.ten_hang,
                    'ma_hang': hang_hoa.ma_hang,
                    'ton_kho': hang_hoa.so_luong_he_thong,
                    'can_xuat': so_luong_xuat
                })

        if not du_ton_kho:
            error_message = "Không đủ tồn kho để xuất các mặt hàng sau: "
            for hang in thieu_hang:
                error_message += f"\n- {hang['ten_hang']} (Mã: {hang['ma_hang']}): Tồn kho {hang['ton_kho']}, Cần xuất {hang['can_xuat']}"
            messages.error(request, error_message)
            return redirect('xuatkho-detail', ma_xuat=ma_xuat)

        # Nếu đủ tồn kho, tiến hành duyệt phiếu và giảm tồn kho
        xuatkho.tinh_trang = 'DA_DUYET'
        xuatkho.save()

        # Giảm tồn kho trong hệ thống và trong từng kho
        for chitiet in xuatkho.chi_tiet_xuat.all():
            hang_hoa = chitiet.ma_hang
            so_luong_xuat = chitiet.so_luong_xuat
            kho = xuatkho.kho

            # Giảm tồn kho trong hệ thống
            hang_hoa.so_luong_he_thong -= so_luong_xuat
            hang_hoa.save()

            # Giảm tồn kho trong kho cụ thể
            ton_kho_theo_kho, created = TonKhoTheoKho.objects.get_or_create(
                hang_hoa=hang_hoa,
                kho=kho,
                defaults={'so_luong': 0}
            )

            if ton_kho_theo_kho.so_luong < so_luong_xuat:
                # Nếu kho này không đủ, tìm kho khác để lấy hàng
                con_thieu = so_luong_xuat - ton_kho_theo_kho.so_luong
                ton_kho_theo_kho.so_luong = 0
                ton_kho_theo_kho.save()

                # Tìm các kho khác có hàng
                cac_kho_khac = TonKhoTheoKho.objects.filter(
                    hang_hoa=hang_hoa,
                    so_luong__gt=0
                ).exclude(kho=kho).order_by('kho__loai_kho')

                for ton_kho_khac in cac_kho_khac:
                    if con_thieu <= 0:
                        break

                    if ton_kho_khac.so_luong >= con_thieu:
                        ton_kho_khac.so_luong -= con_thieu
                        ton_kho_khac.save()
                        con_thieu = 0
                    else:
                        con_thieu -= ton_kho_khac.so_luong
                        ton_kho_khac.so_luong = 0
                        ton_kho_khac.save()
            else:
                ton_kho_theo_kho.so_luong -= so_luong_xuat
                ton_kho_theo_kho.save()

        messages.success(request, f"Phiếu xuất kho {ma_xuat} đã được duyệt và đã giảm hàng tồn kho.")

    return redirect('xuatkho-detail', ma_xuat=ma_xuat)


# Hoàn thành phiếu xuất kho
def xuat_kho(request, ma_xuat):
    xuatkho = get_object_or_404(XuatKho, ma_xuat=ma_xuat)

    if xuatkho.tinh_trang == 'DA_DUYET':
        with transaction.atomic():
            xuatkho.tinh_trang = 'DA_XUAT'
            xuatkho.save()

            # Signal sẽ tự động cập nhật số lượng hàng hóa

            messages.success(request, f"Phiếu xuất kho {ma_xuat} đã được xuất kho thành công.")
    else:
        messages.error(request, "Chỉ có thể xuất kho phiếu đã được duyệt.")

    return redirect('xuatkho-detail', ma_xuat=ma_xuat)


# Từ chối phiếu xuất kho
def tu_choi_xuatkho(request, ma_xuat):
    if not is_quan_ly_kho(request.user):
        messages.error(request, "Bạn không có quyền duyệt phiếu nhập kho.")
        return redirect('nhapkho-detail', ma_nhap=ma_xuat)
    xuatkho = get_object_or_404(XuatKho, ma_xuat=ma_xuat)

    if xuatkho.tinh_trang == 'CHO_DUYET':
        xuatkho.tinh_trang = 'TU_CHOI'
        xuatkho.save()
        messages.success(request, f"Phiếu xuất kho {ma_xuat} đã bị từ chối.")
    else:
        messages.error(request, "Chỉ có thể từ chối phiếu đang ở trạng thái chờ duyệt.")

    return redirect('xuatkho-detail', ma_xuat=ma_xuat)


from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.contrib import messages
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db import transaction
from .models import KiemKe, HangHoa, ChiTietKiemKe
from .forms import KiemKeForm, ChiTietKiemKeFormSet, ChiTietKiemKeFormSetUpdate


@login_required
def kiemke_list(request):
    """Hiển thị danh sách phiếu kiểm kê với các bộ lọc."""
    kiemke_list = KiemKe.objects.all().order_by('-ngay_tao')

    # Xử lý các tham số tìm kiếm
    search_ma_kiemke = request.GET.get('ma_kiemke')
    search_muc_dich = request.GET.get('muc_dich')
    search_tinh_trang = request.GET.get('tinh_trang')
    search_tu_ngay = request.GET.get('tu_ngay')
    search_den_ngay = request.GET.get('den_ngay')

    # Áp dụng các bộ lọc
    if search_ma_kiemke:
        kiemke_list = kiemke_list.filter(ma_kiemke__icontains=search_ma_kiemke)
    if search_muc_dich:
        kiemke_list = kiemke_list.filter(muc_dich__icontains=search_muc_dich)
    if search_tinh_trang:
        kiemke_list = kiemke_list.filter(tinh_trang=search_tinh_trang)
    if search_tu_ngay:
        kiemke_list = kiemke_list.filter(thoi_gian__date__gte=search_tu_ngay)
    if search_den_ngay:
        kiemke_list = kiemke_list.filter(thoi_gian__date__lte=search_den_ngay)

    # Phân trang
    paginator = Paginator(kiemke_list, 10)  # 10 phiếu mỗi trang
    page = request.GET.get('page')
    try:
        kiemke_page = paginator.page(page)
    except PageNotAnInteger:
        kiemke_page = paginator.page(1)
    except EmptyPage:
        kiemke_page = paginator.page(paginator.num_pages)

    context = {
        'kiemke_list': kiemke_page,
        'page_obj': kiemke_page,
        'is_paginated': kiemke_page.has_other_pages(),
        'paginator': paginator,
        'search_params': request.GET.urlencode()
    }
    return render(request, 'cuoiky/kiemke_list.html', context)


@login_required
def kiemke_create(request):
    """Tạo mới phiếu kiểm kê."""
    if request.method == 'POST':
        kiemke_form = KiemKeForm(request.POST)

        # Xử lý dữ liệu chi tiết kiểm kê từ form
        chi_tiet_data = []
        for key, value in request.POST.items():
            if key.startswith('hang_hoa_'):
                index = key.split('_')[-1]
                hang_hoa_id = value
                so_luong_he_thong = request.POST.get(f'so_luong_he_thong_{index}', 0)
                so_luong_tai_kho = request.POST.get(f'so_luong_tai_kho_{index}', 0)
                xu_ly_value = request.POST.get(f'xu_ly_value_{index}', '')  # Lấy giá trị từ hidden field
                kho_id = request.POST.get(f'kho_{index}', None)

                chi_tiet_data.append({
                    'hang_hoa': hang_hoa_id,
                    'kho': kho_id,
                    'so_luong_he_thong': so_luong_he_thong,
                    'so_luong_tai_kho': so_luong_tai_kho,
                    'xu_ly': xu_ly_value  # Sử dụng giá trị từ hidden field
                })

        if kiemke_form.is_valid() and chi_tiet_data:
            try:
                with transaction.atomic():
                    # Lưu phiếu kiểm kê
                    kiemke = kiemke_form.save(commit=False)
                    kiemke.tao_boi = request.user
                    kiemke.tinh_trang = 'CHO_DUYET'  # Đặt trạng thái mặc định là chờ duyệt
                    kiemke.ngay_tao = timezone.now()  # Đặt thời gian tạo

                    # Tạo mã kiểm kê tự động nếu chưa có
                    if not kiemke.ma_kiemke:
                        prefix = 'KK'
                        last_kiemke = KiemKe.objects.order_by('-ma_kiemke').first()
                        if last_kiemke and last_kiemke.ma_kiemke.startswith(prefix):
                            try:
                                last_number = int(last_kiemke.ma_kiemke.replace(prefix, ''))
                                new_number = last_number + 1
                            except ValueError:
                                new_number = 1
                        else:
                            new_number = 1
                        kiemke.ma_kiemke = f"{prefix}{new_number:04d}"

                    kiemke.save()

                    # Lưu chi tiết kiểm kê
                    for item in chi_tiet_data:
                        hang_hoa = HangHoa.objects.get(pk=item['hang_hoa'])
                        kho = Kho.objects.get(pk=item['kho'])
                        so_luong_he_thong = int(item['so_luong_he_thong'])
                        so_luong_tai_kho = int(item['so_luong_tai_kho'])
                        chenh_lech = so_luong_tai_kho - so_luong_he_thong

                        ChiTietKiemKe.objects.create(
                            phieu_kiem_ke=kiemke,
                            hang_hoa=hang_hoa,
                            kho=kho,
                            so_luong_he_thong=so_luong_he_thong,
                            so_luong_tai_kho=so_luong_tai_kho,
                            chenh_lech=chenh_lech,
                            xu_ly=item['xu_ly']
                        )

                messages.success(request, 'Phiếu kiểm kê đã được tạo thành công.')
                return redirect('kiemke-detail', ma_kiemke=kiemke.ma_kiemke)
            except Exception as e:
                messages.error(request, f'Lỗi khi tạo phiếu kiểm kê: {str(e)}')
        else:
            if not chi_tiet_data:
                messages.error(request, 'Vui lòng thêm ít nhất một hàng hóa vào phiếu kiểm kê.')
            else:
                messages.error(request, 'Vui lòng kiểm tra lại thông tin phiếu kiểm kê.')
    else:
        # Tạo mã kiểm kê tự động
        prefix = 'KK'
        last_kiemke = KiemKe.objects.order_by('-ma_kiemke').first()
        if last_kiemke and last_kiemke.ma_kiemke.startswith(prefix):
            try:
                last_number = int(last_kiemke.ma_kiemke.replace(prefix, ''))
                new_number = last_number + 1
            except ValueError:
                new_number = 1
        else:
            new_number = 1
        ma_kiemke = f"{prefix}{new_number:04d}"

        kiemke_form = KiemKeForm(initial={
            'ma_kiemke': ma_kiemke,
            'tinh_trang': 'CHO_DUYET'  # Đặt trạng thái mặc định là chờ duyệt
        })

    context = {
        'kiemke_form': kiemke_form,
        'all_hanghoa': HangHoa.objects.all(),
        'all_kho': Kho.objects.all(),
        'ton_kho_theo_kho': TonKhoTheoKho.objects.all(),
        'now': timezone.now()
    }
    return render(request, 'cuoiky/kiemke_create.html', context)


@login_required
def kiemke_detail(request, ma_kiemke):
    """Xem chi tiết phiếu kiểm kê."""
    kiemke = get_object_or_404(KiemKe, ma_kiemke=ma_kiemke)
    context = {
        'kiemke': kiemke,
    }
    return render(request, 'cuoiky/kiemke_detail.html', context)


@login_required
def kiemke_update(request, ma_kiemke):
    """Cập nhật phiếu kiểm kê."""
    kiemke = get_object_or_404(KiemKe, ma_kiemke=ma_kiemke)

    # Không cho phép cập nhật phiếu đã duyệt
    if kiemke.tinh_trang != 'CHO_DUYET':
        messages.error(request, 'Không thể cập nhật phiếu kiểm kê đã được duyệt hoặc từ chối.')
        return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

    if request.method == 'POST':
        kiemke_form = KiemKeForm(request.POST, instance=kiemke)

        # Xử lý dữ liệu chi tiết kiểm kê từ form
        chi_tiet_data = []
        chi_tiet_ids = []

        for key, value in request.POST.items():
            if key.startswith('hang_hoa_'):
                index = key.split('_')[-1]
                hang_hoa_id = value
                chi_tiet_id = request.POST.get(f'chi_tiet_id_{index}', '')
                so_luong_he_thong = request.POST.get(f'so_luong_he_thong_{index}', 0)
                so_luong_tai_kho = request.POST.get(f'so_luong_tai_kho_{index}', 0)
                xu_ly_value = request.POST.get(f'xu_ly_value_{index}', '')  # Lấy giá trị từ hidden field
                delete = request.POST.get(f'delete_{index}', '') == 'on'

                if not delete:
                    chi_tiet_data.append({
                        'id': chi_tiet_id,
                        'hang_hoa': hang_hoa_id,
                        'so_luong_he_thong': so_luong_he_thong,
                        'so_luong_tai_kho': so_luong_tai_kho,
                        'xu_ly': xu_ly_value  # Sử dụng giá trị từ hidden field
                    })
                    if chi_tiet_id:
                        chi_tiet_ids.append(int(chi_tiet_id))

        if kiemke_form.is_valid() and chi_tiet_data:
            try:
                with transaction.atomic():
                    # Lưu phiếu kiểm kê
                    kiemke = kiemke_form.save(commit=False)
                    kiemke.tinh_trang = 'CHO_DUYET'  # Đảm bảo trạng thái vẫn là chờ duyệt
                    kiemke.save()

                    # Xóa chi tiết không còn trong form
                    kiemke.chi_tiet_kiem_ke.exclude(id__in=chi_tiet_ids).delete()

                    # Cập nhật hoặc tạo mới chi tiết kiểm kê
                    for item in chi_tiet_data:
                        hang_hoa = HangHoa.objects.get(pk=item['hang_hoa'])
                        so_luong_he_thong = int(item['so_luong_he_thong'])
                        so_luong_tai_kho = int(item['so_luong_tai_kho'])
                        chenh_lech = so_luong_tai_kho - so_luong_he_thong

                        if item['id']:
                            # Cập nhật chi tiết hiện có
                            chi_tiet = ChiTietKiemKe.objects.get(id=item['id'])
                            chi_tiet.hang_hoa = hang_hoa
                            chi_tiet.so_luong_he_thong = so_luong_he_thong
                            chi_tiet.so_luong_tai_kho = so_luong_tai_kho
                            chi_tiet.chenh_lech = chenh_lech
                            chi_tiet.xu_ly = item['xu_ly']
                            chi_tiet.save()
                        else:
                            # Tạo chi tiết mới
                            ChiTietKiemKe.objects.create(
                                phieu_kiem_ke=kiemke,
                                hang_hoa=hang_hoa,
                                so_luong_he_thong=so_luong_he_thong,
                                so_luong_tai_kho=so_luong_tai_kho,
                                chenh_lech=chenh_lech,
                                xu_ly=item['xu_ly']
                            )

                messages.success(request, 'Phiếu kiểm kê đã được cập nhật thành công.')
                return redirect('kiemke-detail', ma_kiemke=kiemke.ma_kiemke)
            except Exception as e:
                messages.error(request, f'Lỗi khi cập nhật phiếu kiểm kê: {str(e)}')
        else:
            if not chi_tiet_data:
                messages.error(request, 'Vui lòng thêm ít nhất một hàng hóa vào phiếu kiểm kê.')
            else:
                messages.error(request, 'Vui lòng kiểm tra lại thông tin phiếu kiểm kê.')
    else:
        kiemke_form = KiemKeForm(instance=kiemke)

    context = {
        'kiemke_form': kiemke_form,
        'kiemke': kiemke,
        'all_hanghoa': HangHoa.objects.all()
    }
    return render(request, 'cuoiky/kiemke_update.html', context)


@login_required
def kiemke_delete(request, ma_kiemke):
    if not is_quan_ly_kho(request.user):
        messages.error(request, 'Bạn không có  quyền xóa phiếu kiểm kê.')
        return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

    """Xóa phiếu kiểm kê."""
    kiemke = get_object_or_404(KiemKe, ma_kiemke=ma_kiemke)

    # Không cho phép xóa phiếu đã duyệt
    if kiemke.tinh_trang != 'CHO_DUYET':
        messages.error(request, 'Không thể xóa phiếu kiểm kê đã được duyệt hoặc từ chối.')
        return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # Xóa chi tiết kiểm kê trước
                kiemke.chi_tiet_kiem_ke.all().delete()
                # Xóa phiếu kiểm kê
                kiemke.delete()

            messages.success(request, f'Phiếu kiểm kê {ma_kiemke} đã được xóa thành công.')
            return redirect('kiemke-list')
        except Exception as e:
            messages.error(request, f'Lỗi khi xóa phiếu kiểm kê: {str(e)}')
            return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

    # Nếu không phải POST, chuyển hướng về trang danh sách
    return redirect('kiemke-list')


@login_required
def duyet_kiemke(request, ma_kiemke):
    if not is_quan_ly_kho(request.user):
        messages.error(request, "Bạn không có quyền duyệt phiếu kiểm kê.")
        return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

    kiemke = get_object_or_404(KiemKe, ma_kiemke=ma_kiemke)

    if kiemke.tinh_trang != 'CHO_DUYET':
        messages.error(request, 'Chỉ có thể duyệt phiếu đang ở trạng thái chờ duyệt.')
        return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

    try:
        with transaction.atomic():
            # Cập nhật trạng thái phiếu
            kiemke.tinh_trang = 'DA_DUYET'
            kiemke.save()

        messages.success(request, f'Phiếu kiểm kê {ma_kiemke} đã được duyệt .')
    except Exception as e:
        messages.error(request, f'Lỗi khi duyệt phiếu kiểm kê: {str(e)}')

    return redirect('kiemke-detail', ma_kiemke=ma_kiemke)


@login_required
def tu_choi_kiemke(request, ma_kiemke):
    if not is_quan_ly_kho(request.user):
        messages.error(request, "Bạn không có quyền từ chối phiếu kiểm kê.")
        return redirect('kiemke-detail', ma_kiemke=ma_kiemke)
    """Từ chối phiếu kiểm kê."""
    kiemke = get_object_or_404(KiemKe, ma_kiemke=ma_kiemke)

    if kiemke.tinh_trang != 'CHO_DUYET':
        messages.error(request, 'Chỉ có thể từ chối phiếu đang ở trạng thái chờ duyệt.')
        return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

    kiemke.tinh_trang = 'TU_CHOI'
    kiemke.save()

    messages.success(request, f'Phiếu kiểm kê {ma_kiemke} đã bị từ chối.')
    return redirect('kiemke-detail', ma_kiemke=ma_kiemke)

# cuoiky/views.py

from django.views.generic import ListView
from .models import KhachHang, NhaCungCap
from django.db.models import Q

class KhachHangListView(ListView):
    model = KhachHang
    template_name = 'cuoiky/khachhang_list.html'
    context_object_name = 'khachhang_list'

    def get_queryset(self):
        qs = super().get_queryset()
        q = self.request.GET.get('q')
        if q:
            qs = qs.filter(
                Q(ma_khachhang__icontains=q) |
                Q(ten_khachhang__icontains=q) |
                Q(sdt__icontains=q)
            )
        return qs


class NhaCungCapListView(ListView):
    model = NhaCungCap
    template_name = 'cuoiky/nhacungcap_list.html'
    context_object_name = 'nhacungcap_list'

    def get_queryset(self):
        qs = super().get_queryset()
        q = self.request.GET.get('q')
        if q:
            qs = qs.filter(
                Q(ma_ncc__icontains=q) |
                Q(ten_ncc__icontains=q) |
                Q(sdt__icontains=q)
            )
        return qs


class NhaCungCapDetailView(DetailView):
    model = NhaCungCap
    template_name = 'cuoiky/nhacungcap_detail.html'
    context_object_name = 'ncc'

class NhaCungCapUpdateView(UpdateView):
    model = NhaCungCap
    form_class = NhaCungCapForm
    template_name = 'cuoiky/nhacungcap_form.html'
    success_url = reverse_lazy('nhacungcap-list')

    def dispatch(self, request, *args, **kwargs):
        if not is_quan_ly_kho(request.user):
            messages.error(request, "Bạn không có quyền sửa thông tin nhà cung cấp.")
            return redirect('nhacungcap-detail', pk=kwargs.get('pk'))
        return super().dispatch(request, *args, **kwargs)

class NhaCungCapCreateView(CreateView):
    model = NhaCungCap
    form_class = NhaCungCapForm
    template_name = 'cuoiky/nhacungcap_form.html'
    success_url = reverse_lazy('nhacungcap-list')
    def dispatch(self, request, *args, **kwargs):
        if not is_quan_ly_kho(request.user):
            messages.error(request, "Bạn không có quyền tạo thông tin nhà cung cấp.")
            return redirect('nhacungcap-list')
        return super().dispatch(request, *args, **kwargs)

    def get_initial(self):
        initial = super().get_initial()
        last_item = NhaCungCap.objects.order_by('-ma_ncc').first()
        if last_item and last_item.ma_ncc.startswith("NCC"):
            try:
                last_number = int(last_item.ma_ncc.replace("NCC", ""))
                new_number = last_number + 1
            except ValueError:
                new_number = 1
        else:
            new_number = 1
        initial['ma_ncc'] = f"NCC{new_number:04d}"
        return initial

@login_required
def delete_nhacungcap(request, pk):
    if not is_quan_ly_kho(request.user):
        messages.error(request, "Bạn không có quyền xóa nhà cung cấp.")
        return redirect('nhacungcap-detail', pk=pk)

    ncc = get_object_or_404(NhaCungCap, pk=pk)
    try:
        ncc.delete()
        messages.success(request, f'Đã xóa nhà cung cấp: {ncc.ten_ncc}')
    except Exception as e:
        messages.error(request, f'Không thể xóa: {e}')
    return redirect('nhacungcap-list')


class KhoListView(ListView):
    model = Kho
    template_name = 'cuoiky/kho_list.html'
    context_object_name = 'kho_list'

class KhoDetailView(DetailView):
    model = Kho
    template_name = 'cuoiky/kho_detail.html'
    context_object_name = 'kho'
    pk_url_kwarg = 'ma_kho'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['ton_kho'] = TonKhoTheoKho.objects.filter(kho=self.object).select_related('hang_hoa')
        return context

